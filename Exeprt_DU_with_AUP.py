# Импортируем модули для создания интерфейса и работы с Word-документами
import tkinter as tk
from tkinter import messagebox
from docx import Document

# Функция для извлечения нужных разделов из исходного .docx по ключевым словам
def extract_sections(doc, keywords):
    result = []              # Список для хранения найденных параграфов
    capture = False          # Флаг захвата текста
    current_keyword = None   # Текущий активный раздел

    # Перебираем все параграфы документа
    for para in doc.paragraphs:
        text = para.text.strip()  # Удаляем пробелы


        # Проверяем начало и конец каждого раздела по ключевым словам
        for keyword in keywords:
            if keyword in text and not text.endswith("конец"):
                capture = True
                current_keyword = keyword
                result.append((para, keyword))  # Сохраняем параграф
                break
            elif keyword in text and text.endswith("конец") and capture and current_keyword == keyword:
                result.append((para, keyword))  # Добавляем завершающий параграф
                capture = False
                current_keyword = None
                break
        else:
            if capture:
                result.append((para, current_keyword))  # Добавляем текст внутри секции

    return result  # Возвращаем список выбранных параграфов

# Функция создания нового документа из выбранных разделов
def generate_doc():
    try:
        src_doc = Document("исходный.docx")  # Загружаем исходный документ
        dst_doc = Document()                 # Создаем новый пустой документ


        # ПРАВИТЬ
        selected_keywords = []              # Список выбранных разделов
        if var_a1.get():
            selected_keywords.append("a1")
        if var_a3.get():
            selected_keywords.append("a3")
        if var_a4.get():
            selected_keywords.append("a4")
        if var_a5.get():
            selected_keywords.append("a5")
        if var_b1.get():
            selected_keywords.append("b1")
        if var_b4.get():
            selected_keywords.append("b4")
        if var_b5.get():
            selected_keywords.append("b5")
        if var_b6.get():
            selected_keywords.append("b6")
        if var_b7.get():
            selected_keywords.append("b7")
        if var_b8.get():
            selected_keywords.append("b8")
        if var_b9.get():
            selected_keywords.append("b9")
        if var_c1.get():
            selected_keywords.append("c1")
        if var_c2.get():
            selected_keywords.append("c2")
        if var_c3.get():
            selected_keywords.append("c3")
        if var_c4.get():
            selected_keywords.append("c4")
        if var_c5.get():
            selected_keywords.append("c5")
        if var_c6.get():
            selected_keywords.append("c6")
        if var_c7.get():
            selected_keywords.append("c7")
        if var_c8.get():
            selected_keywords.append("c8")
        if var_c9.get():
            selected_keywords.append("c9")
        if var_d1.get():
            selected_keywords.append("d1")
        if var_d2.get():
            selected_keywords.append("d2")
        if var_d3.get():
            selected_keywords.append("d3")
        if not selected_keywords:


            # Если ничего не выбрано — предупреждение
            messagebox.showwarning("Внимание", "Выберите хотя бы один раздел.")
            return

        # Извлекаем текст из выбранных разделов
        sections = extract_sections(src_doc, selected_keywords)

        # Копируем текст и стили из исходного в новый документ
        for para, _ in sections:
            new_para = dst_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color else None

        dst_doc.save("готовый.docx")  # Сохраняем итоговый файл
        messagebox.showinfo("Готово", "Файл 'готовый.docx' успешно создан.")
    except Exception as e:
        # Показываем сообщение об ошибке, если что-то пошло не так
        messagebox.showerror("Ошибка", str(e))


# === Создание интерфейса ===
root = tk.Tk()                    # Главное окно
root.title("Exert")  # Заголовок окна
root.geometry("600x500")         # Размер окна

# Верхняя рамка с прокруткой
top_frame = tk.Frame(root)
top_frame.pack(fill="both", expand=True)

canvas = tk.Canvas(top_frame)    # Область прокрутки
scrollbar = tk.Scrollbar(top_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

# Обновляем область прокрутки при изменении содержимого
scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

# Встраиваем прокручиваемый фрейм внутрь канваса
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)  # Размещение канваса
scrollbar.pack(side="right", fill="y")              # Размещение ползунка

# ПРАВИТЬ
# Переменные для чекбоксов
var_a1 = tk.BooleanVar()
var_a3 = tk.BooleanVar()
var_a4 = tk.BooleanVar()
var_a5 = tk.BooleanVar()
var_b1 = tk.BooleanVar()
var_b4 = tk.BooleanVar()
var_b5 = tk.BooleanVar()
var_b6 = tk.BooleanVar()
var_b7 = tk.BooleanVar()
var_b8 = tk.BooleanVar()
var_b9 = tk.BooleanVar()
var_c1 = tk.BooleanVar()
var_c2 = tk.BooleanVar()
var_c3 = tk.BooleanVar()
var_c4 = tk.BooleanVar()
var_c5 = tk.BooleanVar()
var_c6 = tk.BooleanVar()
var_c7 = tk.BooleanVar()
var_c8 = tk.BooleanVar()
var_c9 = tk.BooleanVar()
var_d1 = tk.BooleanVar()
var_d2 = tk.BooleanVar()
var_d3 = tk.BooleanVar()

# ПРАВИТЬ
# Подпись и чекбоксы для выбора разделов
tk.Label(scrollable_frame, text="Выберите необходимые параметры:").pack(anchor="w", pady=(5, 5))
tk.Checkbutton(scrollable_frame, text="1. Атриум и/или пассаж", variable=var_a1).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="3. Общественные Помещения более 50м2/ больше 200м.кв", variable=var_a3).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="4. Гардеробные (раздевалки) более 200м2", variable=var_a4).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="5. Производственные и складские помещения категорий А, Б, В1, В2, В3 \nв зданиях I - IV степени огнестойкости, а также В4 или Г площадью 50 м2/ и более \nв зданиях IV степени огнестойкости больше 200м.кв", variable=var_a5).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="10. Производственные и складские помещения категорий А, Б, В1, В2, В3 \nв зданиях I - IV степени огнестойкости, а также В4 или Г площадью 50 м2/ и более \nв зданиях IV степени огнестойкости высокостелажка более 200м.кв", variable=var_b1).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="13. Коридор цокольных (заглубленных менее 0,5 м) и наземных этажей \nжилых, общественных, административно-бытовых и многофункциональных зданий \nвысотой более 28 м", variable=var_b4).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="14. Коридор подвал и цоколь (заглубленный более 0,5м) во всех зданиях", variable=var_b5).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="""15. Коридор цокольных (заглубленных менее 0,5 м) и наземных этажей \nбез естественного проветривания при пожаре длиной более 15 м \nв зданиях с числом этажей два и более:
- производственных и складских категорий А, Б, В;
- общественных и административно-бытовых;
- многофункциональных""", variable=var_b6).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="16. Объект культурного наследия СП388", variable=var_b7).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="17. Шахты лифтов (при отсутствии у выходов из них тамбур-шлюзов, \nзащищаемых приточной противодымной вентиляцией),  \nв зданиях с незадымляемыми лестничными клетками", variable=var_b8).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="18. Лифт для пожарных", variable=var_b9).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="19. Незадымляемая ЛК типа Н2", variable=var_c1).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="20. Тамбур шлюз перед незадымляемой ЛК Н3", variable=var_c2).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="21. Тамбур шлюз перед лифтом в подземной автостоянке", variable=var_c3).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="22. Тамбур-шлюз при внутренних технологических лестницах, \nпредназначенных для сообщения между подвальным этажом \nи цокольным или первым этажом", variable=var_c4).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="23. Тамбур шлюз на входах из коридоров в атриумы и пассажи \nс уровней подземных, подвальных и цокольных этажей", variable=var_c5).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="24.	Тамбур шлюз перед незадымляемой лестничной клеткой типа Н2 \nв многофункциональных зданиях и комплексах высотой более 28 м, \nв жилых зданиях высотой более 75 м, в общественных зданиях высотой более 50 м", variable=var_c6).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="25. Тамбур шлюз отделяющий помещения для хранения автомобилей \nзакрытых надземных и подземных автостоянок от помещений общественного назначения", variable=var_c7).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="26. Тамбур шлюз отделяющий помещения для хранения автомобилей \nот изолированных рамп подземных автостоянок", variable=var_c8).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="27. Тамбур шлюз при выходах в вестибюли из незадымляемых лестничных клеток типа Н2", variable=var_c9).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="28. Тамбур-шлюз (лифтовой холл) при выходе из лифта в подвальные, подземные этажи", variable=var_d1).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="29. Зона безопасности МГН", variable=var_d2).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="30. Помещения, сообщающиеся с помещениями (тамбурами, ЛК), \nв которых предусмотрен подпор", variable=var_d3).pack(anchor="w")

# Нижняя рамка — кнопка запуска
bottom_frame = tk.Frame(root)
bottom_frame.pack(fill="x", pady=10)

# Кнопка создания нового документа
tk.Button(bottom_frame, text="Создать готовый.docx", command=generate_doc,
          height=2, font=("Arial", 11, "bold")).pack()

# Запуск графического интерфейса
root.mainloop()

# Напоминание:
# Чтобы встроить иконку в .exe, используй:

# pyinstaller --onefile --noconsole --icon=icon.ico main.py
